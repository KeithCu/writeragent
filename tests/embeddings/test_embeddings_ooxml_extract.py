# WriterAgent - AI Writing Assistant for LibreOffice
# Copyright (c) 2026 KeithCu (modifications and relicensing)
#
# SPDX-License-Identifier: GPL-3.0-or-later
"""Tests for plugin.embeddings.venv.embeddings_ooxml_extract."""

from __future__ import annotations

import zipfile
from pathlib import Path
from unittest.mock import patch

import pytest

from plugin.embeddings.venv import embeddings_ooxml_extract as ooxml


def test_extract_csv_rows(tmp_path: Path):
    path = tmp_path / "data.csv"
    path.write_text("a,b\n1,2\n\n", encoding="utf-8")
    assert ooxml.extract_csv_rows(str(path)) == ["a\tb", "1\t2"]


def test_extract_plaintext_paragraphs_blank_lines(tmp_path: Path):
    path = tmp_path / "notes.txt"
    path.write_text("First block\n\nSecond block\n", encoding="utf-8")
    assert ooxml.extract_plaintext_paragraphs(str(path)) == ["First block", "Second block"]


def test_extract_plaintext_paragraphs_lines(tmp_path: Path):
    path = tmp_path / "lines.txt"
    path.write_text("alpha\nbeta\n", encoding="utf-8")
    assert ooxml.extract_plaintext_paragraphs(str(path)) == ["alpha", "beta"]


def test_extract_rtf_paragraphs(tmp_path: Path):
    path = tmp_path / "doc.rtf"
    path.write_text(r"{\rtf1 hello \par world}", encoding="utf-8")
    assert ooxml.extract_rtf_paragraphs(str(path)) == ["hello", "world"]


def test_extract_pptx_passages(tmp_path: Path):
    slide_xml = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
 xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:cSld><p:spTree><p:sp><p:txBody><a:p><a:r><a:t>Slide text</a:t></a:r></a:p></p:txBody></p:sp></p:spTree></p:cSld>
</p:sld>"""
    pptx = tmp_path / "deck.pptx"
    with zipfile.ZipFile(pptx, "w") as zf:
        zf.writestr("ppt/slides/slide1.xml", slide_xml)
    passages = ooxml.extract_pptx_passages(str(pptx))
    assert passages == ["[Slide: Slide1]\tSlide text"]


def test_extract_docx_paragraphs_uses_python_docx(tmp_path: Path):
    pytest.importorskip("docx")
    from docx import Document

    path = tmp_path / "file.docx"
    document = Document()
    document.add_paragraph("Hello")
    document.add_paragraph("")
    document.add_paragraph("World")
    document.save(path)
    assert ooxml.extract_docx_paragraphs(str(path)) == ["Hello", "World"]


def test_extract_spreadsheet_rows_xlsx(tmp_path: Path):
    pytest.importorskip("openpyxl")
    from openpyxl import Workbook

    path = tmp_path / "book.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Budget"
    ws["A1"] = "Revenue"
    ws["B1"] = 100
    wb.save(path)
    rows = ooxml.extract_spreadsheet_rows(str(path))
    assert rows == ["[Sheet: Budget]\tRevenue\t100"]
